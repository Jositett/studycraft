"""
StudyCraft – DOCX export.

Converts Markdown guide text into a styled .docx file using python-docx.
"""

from __future__ import annotations

import re
from pathlib import Path

from rich.console import Console

console = Console()


def _hex_to_rgb(h: str) -> tuple[int, int, int]:
    h = h.lstrip("#")
    if len(h) != 6:
        return (0, 0, 0)
    return int(h[0:2], 16), int(h[2:4], 16), int(h[4:6], 16)


def export_docx(
    full_markdown: str, output_path: Path, theme: "object | None" = None
) -> Path:
    """Export Markdown text to a styled DOCX file."""
    from docx import Document  # type: ignore
    from docx.shared import Pt, Inches, RGBColor  # type: ignore
    from docx.oxml.ns import qn  # type: ignore
    from docx.oxml import OxmlElement  # type: ignore

    h1c = RGBColor(*_hex_to_rgb(theme.h1)) if theme else RGBColor(37, 99, 235)
    h2c = RGBColor(*_hex_to_rgb(theme.h2)) if theme else RGBColor(30, 58, 138)
    qc = RGBColor(*_hex_to_rgb(theme.quote_fg)) if theme else RGBColor(30, 64, 175)

    doc = Document()

    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(11)

    # TOC field
    doc.add_heading("Table of Contents", level=1)
    p = doc.add_paragraph()
    run = p.add_run()
    fld_char_begin = OxmlElement("w:fldChar")
    fld_char_begin.set(qn("w:fldCharType"), "begin")
    run._r.append(fld_char_begin)
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instr)
    fld_char_end = OxmlElement("w:fldChar")
    fld_char_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char_end)
    doc.add_page_break()

    for line in full_markdown.splitlines():
        stripped = line.strip()

        if not stripped or stripped == "---":
            continue

        if stripped.startswith("# "):
            p = doc.add_heading(stripped[2:], level=1)
            for run in p.runs:
                run.font.color.rgb = h1c
        elif stripped.startswith("## "):
            p = doc.add_heading(stripped[3:], level=2)
            for run in p.runs:
                run.font.color.rgb = h2c
        elif stripped.startswith("### "):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith("| ") and "---" not in stripped:
            doc.add_paragraph(stripped, style="List Bullet")
        elif stripped.startswith("- ") or stripped.startswith("* "):
            doc.add_paragraph(stripped[2:], style="List Bullet")
        elif re.match(r"^\d+\.\s", stripped):
            doc.add_paragraph(stripped, style="List Number")
        elif stripped.startswith("> "):
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.5)
            run = p.add_run(stripped[2:])
            run.italic = True
            run.font.color.rgb = qc
        elif stripped.startswith("```"):
            continue
        elif stripped.startswith("**") and stripped.endswith("**"):
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.bold = True
        else:
            doc.add_paragraph(stripped)

    doc.save(str(output_path))
    console.print(f"[green]\u2713 DOCX[/green]     \u2192 {output_path}")
    return output_path
